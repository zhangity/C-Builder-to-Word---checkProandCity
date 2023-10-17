# encoding: utf-8
#安装新库超时的方法 pip install --index https://pypi.mirrors.ustc.edu.cn/simple cpca
import cpca
import docx

#阿宝在这里修改文件的路径，冒号后面要多一个\
word_path = "C:\\Users\Zhangity\Desktop\凤桥安置房施组.docx"
def printHello():
    print("Hello World!")
# 1.界面读取和C++联动 上传自己的python工程和该过的源码到github 2.支持pdf  3.剔除输入的标的
def getcityname(location_str):
    #location_str = ["衢州开化县","徐汇区虹漕路461号58号楼5楼", "泉州市洛江区万安塘西工业区", "北京朝阳区北苑华贸城"]
    #location_str = ["衢州开化县","","安塘西工业区以","徐汇区虹漕路461号58号楼5楼"]
    df = cpca.transform_notnone(location_str)
    print(df)

# 提取一段话中的省市信息
word_path = word_path.replace("\\", "/")
doc = docx.Document(word_path)
if __name__ == "__main__":
    # 提取段落的文字（不含表格内文字）
    paragraph_data = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() != "":
            paragraph_data.append(paragraph.text.strip())
    #getcityname(paragraph_data)

    # 提取表格内文字
    tables_data = []
    for table in doc.tables:
        for row in table.rows:
            cell_data = []
            for cell in row.cells:
                if cell.text.strip() != "":
                    tables_data.append(cell.text.strip())
    #getcityname(tables_data)

    #合并段落和表格的所有文字
    all_datas = paragraph_data + tables_data
    getcityname(all_datas)