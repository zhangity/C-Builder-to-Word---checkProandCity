# encoding: utf-8
# python -m pip install --upgrade pip

# 提取一段话中的省市信息
#安装新库超时的方法 pip install --index https://pypi.mirrors.ustc.edu.cn/simple cpca python-docx
# 升级 Microsoft Visual C++ 报错参考 https://wenku.csdn.net/answer/6gfuez0wd6和https://blog.csdn.net/weixin_43955488/article/details/106739251 只用复制2个文件即可

import cpca
import docx
import datetime
import concurrent.futures
import threading
import chardet

# 阿宝在这里修改文件的路径，冒号后面要多一个\
# word_path = "C:\\Users\Zhangity\Desktop\凤桥安置房施组.docx"  #凤桥安置房施组.docx
# word_path = word_path.replace("\\", "/")
# doc = docx.Document(word_path)

def testCPlus(a):
    s = list()
    s.append("张雨亭")
    s.append(a)
    s.append("率")
    ss = set()
    # s.append(str(bytes(a, 'utf-8')))
    try:
        s.append("try1")
        res = a.encode('utf-8')
        s.append(res)
        s.append("try2")
    except Exception as e:
        s.append("error")
        s.append(e)
    finally:
        s.append("finally")
    print(s)
    return s

# 1.界面读取和C++联动 上传自己的python工程和该过的源码到github 2.支持pdf  3.剔除输入的标的
def getcityname(location_str):
    #location_str = ["衢州开化县","徐汇区虹漕路461号58号楼5楼", "泉州市洛江区万安塘西工业区", "北京朝阳区北苑华贸城"]
    #location_str = ["衢州开化县","","安塘西工业区以","徐汇区虹漕路461号58号楼5楼"]
    # location_str = ['EPC项目管理宗旨：工程总承包（EPC）采用科学的项目管理技术和项目管理方法，进行工程质量、安全、成本、进度等方面的综合控制，实行项目经理责任制和项目成本核算制。项目经理责任制和项目成本核算制是实施工程项目管理的核心内容，是实现工程总承包（EPC）项目管理的关键；项目管理人员按照工程建设的有关法律、法规、技术规范的要求，用系统工程的理论、观点和方法，进行有效的规划、决策、组织、协调、控制等系统性的、科学的管理，根据已签订的工程项目管理合同和其他合同性文件、相关法律、行政法规、以及业主现已完成的各项前期工作，调动各方面资源，代表或协助业主对项目前期管理、工程施工阶段的管理、竣工移交阶段进行全过程的工程项目总控制。为保证项目产品和服务的质量，满足合同及相关方的要求，公司建立覆盖设计、采购、施工、试运行全过程的项目管理体系、质量管理体系、HSE （职业健康、安全管理和环境管理）体系。']
    df = cpca.transform_notnone(location_str)
    # print(df)
    # 把set转成list
    return list(df)

#拆分list l 是要拆分的list，n是要拆分多少个子list 返回值就是一个嵌套List
def divide_chunks(l, n):
    # looping till length l
    for i in range(0, len(l), n):
        yield l[i:i + n]

def checkProCity(sWordFilePath):
    # 提取段落的文字（不含表格内文字） 先转换路径的斜杠方向，再读取word文件
    sWordFilePath = sWordFilePath.replace("\\", "/")
    doc = docx.Document(sWordFilePath)
    paragraph_data = []
    time6 = datetime.datetime.now()
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() != "":
            paragraph_data.append(paragraph.text.strip())
            # if len(paragraph_data) >= 1:
                # break
    # getcityname(paragraph_data)
    time7 = datetime.datetime.now()
    print("段落提取耗时：", time7 - time6)
    # 提取表格内文字
    tables_data = []
    for table in doc.tables:
        for row in table.columns:
            cell_data = []
            for cell in row.cells:
                if cell.text.strip() != "":
                    tables_data.append(cell.text.strip())
    # getcityname(tables_data)
    time8 = datetime.datetime.now()
    print("表格提取耗时：", time8 - time7)
    # 合并段落和表格的所有文字
    all_datas = paragraph_data + tables_data
    # return getcityname(all_datas)

    # n = len(all_datas) % 5
    # sub_lists = [tables_data[i:i + n] for i in range(0, len(tables_data), n)]
    sub_lists = list(divide_chunks(all_datas, 7))
    # print("sub_lists chahgndu " + str(len(paragraph_data)))
    # print("sub_lists chahgndu " + str(len(tables_data)))
    # # print("sub_lists chahgndu " + str(len(all_datas)))
    # all_datas = all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas + all_datas
    print("sub_lists chahgndu " + str(len(sub_lists)))
    result = checkMulti(sub_lists)
    print(set(result))
    return list(set(result))
    # return getcityname(tables_data)
    # return parallel_process_file(sub_lists)

#加入多线程 -----------------------------------
# def getcitynameMulti(input_file_path):
#     df = cpca.transform_notnone(input_file_path)
#     return df
#
# def parallel_process_file(input_file_list, max_workers=8):
#     with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
#         future_to_file = {executor.submit(getcitynameMulti, input_file_path): input_file_path for input_file_path in input_file_list}
#         for future in concurrent.futures.as_completed(future_to_file):
#             input_file_path = future_to_file[future]
#             try:
#                 result = future.result()
#                 print(result)
#             except Exception as e:
#                 print(f'Error processing {input_file_path}: {e}')
def checkMulti(all_datas):
    multi = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=12) as executor:
        ans = [executor.submit(getcityname, i) for i in all_datas]
        for res in concurrent.futures.as_completed(ans):
            # print(res.result())
            if res.result() is not None and len(res.result()) > 0 and multi.count(res.result()) == 0:
                # multi.append(res.result())  这样写返回的是嵌套数组
                multi = multi + res.result()
    return multi

# 主函数，测试入口
if __name__ == "__main__":
    word_path_main = "C:\\Users\Zhangity\Desktop\沈家雅苑0715.docx"  # 凤桥安置房施组.docx 沈家雅苑0715
    time1 = datetime.datetime.now()
    checkProCity(word_path_main)
    # getcityname(None)
    # scheduler = Scheduler()
    # threads_scheduler(4)
    time2 = datetime.datetime.now()
    print("全部耗时：", time2-time1)
    # testCPlus("zyt")